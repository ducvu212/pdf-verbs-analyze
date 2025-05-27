import os
import logging
import sys
from pathlib import Path

# Setup logging first
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s',
    handlers=[logging.StreamHandler(sys.stdout)]
)

logger = logging.getLogger(__name__)

# Import other modules
try:
    import pypdf
    import spacy
    import re
    import time
    import json
    from concurrent.futures import ProcessPoolExecutor, ThreadPoolExecutor
    import hashlib
    from collections import defaultdict
    from fastapi import FastAPI, UploadFile, File, HTTPException, Request, Form, BackgroundTasks
    from fastapi.responses import HTMLResponse, FileResponse
    from fastapi.staticfiles import StaticFiles
    from fastapi.templating import Jinja2Templates
    from fastapi.middleware.cors import CORSMiddleware
    import uvicorn
    from typing import List, Dict, Any
    import nltk
    import pdfplumber
    import tempfile
    import traceback
    from reportlab.lib.pagesizes import letter
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
    from reportlab.lib.styles import getSampleStyleSheet
    from reportlab.lib import colors
    from weasyprint import HTML, CSS
    from datetime import datetime
    
    logger.info("✅ All imports successful")
except ImportError as e:
    logger.error(f"❌ Import failed: {e}")
    sys.exit(1)

# Production configuration
IS_PRODUCTION = os.getenv("RENDER") is not None or os.getenv("RAILWAY_ENVIRONMENT") is not None
PORT = int(os.environ.get("PORT", 8000))

logger.info(f"🔧 Configuration: IS_PRODUCTION={IS_PRODUCTION}, PORT={PORT}")

# Configure paths
if IS_PRODUCTION:
    BASE_DIR = Path("/opt/render/project/src")
else:
    BASE_DIR = Path(__file__).parent

logger.info(f"📁 BASE_DIR: {BASE_DIR}")

UPLOAD_FOLDER = BASE_DIR / "uploads"
EXPORT_FOLDER = BASE_DIR / "exports" 
CACHE_FOLDER = BASE_DIR / "cache"
TEMPLATES_FOLDER = BASE_DIR / "templates"
STATIC_FOLDER = BASE_DIR / "static"

# Ensure directories exist
def ensure_directories():
    """Ensure all required directories exist with proper permissions"""
    directories = [UPLOAD_FOLDER, EXPORT_FOLDER, CACHE_FOLDER, TEMPLATES_FOLDER, STATIC_FOLDER]
    
    for directory in directories:
        try:
            directory.mkdir(parents=True, exist_ok=True)
            os.chmod(directory, 0o755)
            logger.info(f"✅ Directory ensured: {directory}")
            
            # Create .gitkeep for static
            if directory == STATIC_FOLDER:
                gitkeep = directory / ".gitkeep"
                if not gitkeep.exists():
                    gitkeep.write_text("")
                    
        except Exception as e:
            logger.error(f"❌ Failed to create directory {directory}: {e}")

# Initialize directories
ensure_directories()

# Set up logging
if IS_PRODUCTION:
    logging.getLogger().setLevel(logging.INFO)
else:
    logging.getLogger().setLevel(logging.DEBUG)

# Set up FastAPI app
app = FastAPI(
    title="PDF Verb Analyzer",
    description="Analyze verb tenses in PDF documents",
    version="1.0.0"
)

# Add CORS middleware
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# Set up folders and configurations
ALLOWED_EXTENSIONS = {'pdf'}
MAX_WORKERS = 4
CHUNK_SIZE = 100000

# Set up templates and static files
try:
    templates = Jinja2Templates(directory=str(TEMPLATES_FOLDER))
    templates.env.globals.update({"min": min})
    logger.info(f"✅ Templates initialized: {TEMPLATES_FOLDER}")
except Exception as e:
    logger.error(f"❌ Templates initialization failed: {e}")

# Mount static files
try:
    if STATIC_FOLDER.exists():
        app.mount("/static", StaticFiles(directory=str(STATIC_FOLDER)), name="static")
        logger.info(f"✅ Static files mounted: {STATIC_FOLDER}")
    else:
        logger.warning(f"⚠️ Static directory not found: {STATIC_FOLDER}")
except Exception as e:
    logger.error(f"❌ Failed to mount static files: {e}")

# Load spaCy model
try:
    nlp = spacy.load('en_core_web_sm', disable=['ner', 'textcat'])
    logger.info("✅ spaCy model loaded successfully")
except Exception as e:
    logger.error(f"❌ Failed to load spaCy model: {e}")
    nlp = None

# Cache cho kết quả phân tích
analysis_cache = {}

def allowed_file(filename: str) -> bool:
    """Check if the uploaded file has an allowed extension"""
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

def get_cache_path(file_path: str) -> str:
    """Get cache file path based on the input file and its modification time"""
    file_stat = os.stat(file_path)
    file_hash = hashlib.md5(f"{file_path}_{file_stat.st_mtime}".encode()).hexdigest()
    return os.path.join(CACHE_FOLDER, f"{file_hash}.json")

def save_to_cache(cache_path: str, results: Dict[str, Any]) -> None:
    """Save analysis results to cache"""
    with open(cache_path, 'w') as f:
        # Convert tuples to lists for JSON serialization
        serializable_results = json.dumps(results, default=lambda o: list(o) if isinstance(o, tuple) else o)
        f.write(serializable_results)

def load_from_cache(cache_path: str) -> Dict[str, Any]:
    """Load analysis results from cache"""
    with open(cache_path, 'r') as f:
        cached_data = json.loads(f.read())
        # Convert lists back to tuples where needed
        return cached_data

def extract_text_from_pdf(file_path: str) -> str:
    """Extract text from a PDF file using pypdf as primary method, with pdfplumber and OCR as fallbacks"""
    start_time = time.time()
    text = ""
    error_message = ""
    
    # Method 1: Try pypdf (primary method)
    try:
        logger.info("Extracting text using pypdf")
        text = ""
        with open(file_path, 'rb') as file:
            pdf = pypdf.PdfReader(file)
            for page in pdf.pages:
                page_text = page.extract_text() or ""
                text += page_text + "\n"
                
        if text.strip():
            logger.info(f"Successfully extracted {len(text)} characters using pypdf")
            logger.info(f"PDF text extraction completed in {time.time() - start_time:.2f} seconds")
            return text
        else:
            logger.warning("pypdf extracted empty text, trying fallback method")
    except Exception as e:
        error_message = str(e)
        logger.warning(f"pypdf extraction failed: {error_message}")
        logger.warning(traceback.format_exc())
    
    # Method 2: Try pdfplumber as backup method
    try:
        logger.info("Extracting text using pdfplumber (backup method)")
        with pdfplumber.open(file_path) as pdf:
            # Xử lý song song các trang PDF
            with ThreadPoolExecutor(max_workers=MAX_WORKERS) as executor:
                page_texts = list(executor.map(
                    lambda page: page.extract_text() or "",
                    pdf.pages
                ))
            
            text = "\n".join(filter(None, page_texts))
            
            if text.strip():
                logger.info(f"Successfully extracted {len(text)} characters using pdfplumber")
                logger.info(f"PDF text extraction completed in {time.time() - start_time:.2f} seconds")
                return text
            else:
                logger.warning("pdfplumber extracted empty text")
    except Exception as e:
        error_message = f"{error_message}; pdfplumber: {str(e)}"
        logger.warning(f"pdfplumber extraction failed: {str(e)}")
        logger.warning(traceback.format_exc())
    
    # Method 3: OCR using pdf2image and pytesseract as last resort
    if OCR_AVAILABLE:
        try:
            logger.info("Attempting to extract text using OCR (pdf2image + pytesseract)")
            
            # Create a temporary directory for the images
            with tempfile.TemporaryDirectory() as temp_dir:
                logger.info(f"Converting PDF to images in {temp_dir}")
                
                # Convert PDF to images
                images = convert_from_path(
                    file_path, 
                    output_folder=temp_dir,
                    fmt="jpeg",
                    dpi=300  # Higher DPI for better OCR results
                )
                
                # Extract text from each image using OCR
                page_texts = []
                for i, image in enumerate(images):
                    logger.info(f"Performing OCR on page {i+1}/{len(images)}")
                    page_text = pytesseract.image_to_string(image)
                    page_texts.append(page_text)
                
                text = "\n".join(filter(None, page_texts))
                
                if text.strip():
                    logger.info(f"Successfully extracted {len(text)} characters using OCR")
                    logger.info(f"PDF text extraction completed in {time.time() - start_time:.2f} seconds")
                    return text
                else:
                    logger.warning("OCR extracted empty text")
        except Exception as e:
            error_message = f"{error_message}; OCR: {str(e)}"
            logger.warning(f"OCR extraction failed: {str(e)}")
            logger.warning(traceback.format_exc())
    else:
        logger.warning("OCR is not available. skipping OCR extraction.")
    
    # If all methods failed, return whatever we have or raise an exception
    if text.strip():
        logger.warning(f"All primary methods failed, but managed to extract some text ({len(text)} characters)")
        logger.info(f"PDF text extraction completed in {time.time() - start_time:.2f} seconds")
        return text
    else:
        error_msg = f"Failed to extract text from PDF using all available methods. Original errors: {error_message}"
        logger.error(error_msg)
        raise Exception(error_msg)

def split_text_into_chunks(text: str, chunk_size: int = CHUNK_SIZE) -> List[str]:
    """Split text into manageable chunks for processing"""
    # Cố gắng chia theo đoạn văn để giữ nguyên ngữ cảnh
    paragraphs = text.split('\n\n')
    chunks = []
    current_chunk = ""
    
    for para in paragraphs:
        if len(current_chunk) + len(para) <= chunk_size:
            current_chunk += para + "\n\n"
        else:
            if current_chunk:
                chunks.append(current_chunk)
            current_chunk = para + "\n\n"
    
    if current_chunk:  # Thêm đoạn cuối cùng
        chunks.append(current_chunk)
        
    # Nếu không có đoạn văn rõ ràng, chia theo kích thước cố định
    if not chunks:
        return [text[i:i+chunk_size] for i in range(0, len(text), chunk_size)]
        
    return chunks

def process_chunk(chunk: str) -> Dict[str, Any]:
    """Process a single chunk of text and extract verb information"""
    doc = nlp(chunk)
    
    # Initialize counters
    verb_count = 0
    verb_tenses = defaultdict(int)
    passive_count = 0
    passive_tenses = defaultdict(int)
    modal_count = 0
    
    # Create tense group counters
    tense_groups = {
        "simple_present": [],
        "simple_past": [],
        "present_continuous": [],
        "past_continuous": [], 
        "present_perfect": [],
        "past_perfect": [],
        "future_simple": [],
        "modal_verb": [],
        "modal_passive": [],  # Modal + be + V3 (Passive with modals)
        "modal_perfect": [],  # Modal + have + V3 (Perfect aspect with modals)
        "passive_voice": [],
        "other": []
    }
    
    # Passive voice categorization by tense
    passive_voice_by_tense_group = {
        "simple_present": [],  # is/are + V3 (is written)
        "simple_past": [],     # was/were + V3 (was written)
        "present_continuous": [],  # is/are being + V3 (is being written)
        "past_continuous": [],     # was/were being + V3 (was being written)
        "present_perfect": [],     # has/have been + V3 (has been written)
        "past_perfect": [],        # had been + V3 (had been written)
        "future_simple": [],       # will be + V3 (will be written)
        "future_perfect": [],      # will have been + V3 (will have been written)
        "modal_passive": [],       # modal + be + V3 (can be written)
        "other": []                # Other passive forms
    }
    
    # Modal verb categorization by structure
    modal_verb_structures = {
        "modal_base": [],      # Modal + base verb (can go, should work)
        "modal_passive": [],   # Modal + be + V3 (must be cleaned)
        "modal_perfect": [],   # Modal + have + V3 (should have gone)
    }
    
    # Lists to store detailed information
    all_verbs = []
    passive_verbs = []
    modal_verbs = []
    
    # Process each sentence - tối ưu vòng lặp
    for sent in doc.sents:
        # Tạo set các modal words để kiểm tra nhanh hơn
        modal_words = set()
        aux_verbs = {}
        
        # Tạo map lưu trữ thông tin từ để tránh lặp lại kiểm tra
        for token in sent:
            if token.tag_ == "MD":
                modal_words.add(token.i)
            elif token.pos_ == "AUX" and token.lemma_ == "be":
                aux_verbs[token.i] = token
        
        # Phân tích các từ trong câu
        for token in sent:
            if token.pos_ == "VERB" or token.pos_ == "AUX":
                verb_count += 1
                
                # Determine tense
                tense = determine_tense(token)
                verb_tenses[tense] += 1
                all_verbs.append((token.text, tense))
                
                # Add to appropriate tense group
                if tense in tense_groups:
                    tense_groups[tense].append(token.text)
                else:
                    tense_groups["other"].append(token.text)
                
                # Check for passive voice - tối ưu kiểm tra passive voice
                if is_passive(token, aux_verbs):
                    passive_count += 1
                    passive_tenses[tense] += 1
                    passive_verbs.append((token.text, tense))
                    
                    # Categorize passive voice by tense structure
                    sent_text = token.sent.text.lower()
                    
                    # Modal passive: can be written, should be completed, etc.
                    modal_pattern = any(modal in sent_text.split() for modal in 
                                      ["can", "could", "may", "might", "must", "should", "would", "will", "shall"])
                    has_modal_before_be = False
                    
                    for i, word in enumerate(token.sent):
                        if i > 0 and word.lemma_ == "be" and token.sent[i-1].tag_ == "MD":
                            has_modal_before_be = True
                            break
                    
                    if has_modal_before_be:
                        passive_voice_by_tense_group["modal_passive"].append((token.text, tense))
                    # Present Simple Passive: is/are written
                    elif re.search(r'\b(is|are|am)\b', sent_text) and not re.search(r'\b(is|are|am)\s+being\b', sent_text):
                        passive_voice_by_tense_group["simple_present"].append((token.text, tense))
                    # Past Simple Passive: was/were written
                    elif re.search(r'\b(was|were)\b', sent_text) and not re.search(r'\b(was|were)\s+being\b', sent_text):
                        passive_voice_by_tense_group["simple_past"].append((token.text, tense))
                    # Present Continuous Passive: is/are being written
                    elif re.search(r'\b(is|are|am)\s+being\b', sent_text):
                        passive_voice_by_tense_group["present_continuous"].append((token.text, tense))
                    # Past Continuous Passive: was/were being written
                    elif re.search(r'\b(was|were)\s+being\b', sent_text):
                        passive_voice_by_tense_group["past_continuous"].append((token.text, tense))
                    # Present Perfect Passive: has/have been written
                    elif re.search(r'\b(has|have)\s+been\b', sent_text) and not re.search(r'\b(has|have)\s+been\s+being\b', sent_text):
                        passive_voice_by_tense_group["present_perfect"].append((token.text, tense))
                    # Past Perfect Passive: had been written
                    elif re.search(r'\b(had)\s+been\b', sent_text) and not re.search(r'\b(had)\s+been\s+being\b', sent_text):
                        passive_voice_by_tense_group["past_perfect"].append((token.text, tense))
                    # Future Simple Passive: will be written
                    elif re.search(r'\b(will|shall)\s+be\b', sent_text):
                        passive_voice_by_tense_group["future_simple"].append((token.text, tense))
                    # Future Perfect Passive: will have been written
                    elif re.search(r'\b(will|shall)\s+have\s+been\b', sent_text):
                        passive_voice_by_tense_group["future_perfect"].append((token.text, tense))
                    else:
                        passive_voice_by_tense_group["other"].append((token.text, tense))
                
                # Check for modal verbs - cải tiến kiểm tra modal verbs
                if token.tag_ == "MD":
                    modal_count += 1
                    modal_verbs.append(token.text)
                    
                    # Tìm các token liền kề để kiểm tra các mẫu cấu trúc modal verb
                    words_in_sent = list(token.sent)
                    modal_idx = token.i - words_in_sent[0].i  # Vị trí tương đối của modal trong câu
                    
                    # Kiểm tra modal + be + V3 (modal passive)
                    if modal_idx + 2 < len(words_in_sent):  # Đảm bảo có đủ token phía sau
                        next_token = words_in_sent[modal_idx + 1]
                        next_next_token = words_in_sent[modal_idx + 2]
                        
                        if (next_token.lemma_ == "be" and 
                            next_next_token.tag_ == "VBN"):
                            # Đây là cấu trúc Modal + be + V3 (liền kề)
                            modal_verb_structures["modal_passive"].append(
                                (token.text, next_token.text, next_next_token.text)
                            )
                            continue
                    
                    # Kiểm tra modal + have + V3 (modal perfect)
                    if modal_idx + 2 < len(words_in_sent):
                        next_token = words_in_sent[modal_idx + 1]
                        next_next_token = words_in_sent[modal_idx + 2]
                        
                        if (next_token.lemma_ == "have" and 
                            next_next_token.tag_ == "VBN"):
                            # Đây là cấu trúc Modal + have + V3 (liền kề)
                            modal_verb_structures["modal_perfect"].append(
                                (token.text, next_token.text, next_next_token.text)
                            )
                            continue
                    
                    # Kiểm tra modal + base verb (modal base)
                    if modal_idx + 1 < len(words_in_sent):
                        next_token = words_in_sent[modal_idx + 1]
                        
                        if next_token.tag_ == "VB":
                            # Đây là cấu trúc Modal + base verb (liền kề)
                            modal_verb_structures["modal_base"].append(
                                (token.text, next_token.text)
                            )
                            continue
    
    # Trả về kết quả phân tích của đoạn này
    return {
        'verb_count': verb_count,
        'verb_tenses': dict(verb_tenses),
        'tense_groups': tense_groups,
        'passive_count': passive_count,
        'passive_tenses': dict(passive_tenses),
        'passive_voice_by_tense_group': passive_voice_by_tense_group,
        'modal_count': modal_count,
        'modal_verb_structures': modal_verb_structures,
        'all_verbs': all_verbs,
        'passive_verbs': passive_verbs,
        'modal_verbs': modal_verbs
    }

def merge_chunk_results(results_list: List[Dict[str, Any]]) -> Dict[str, Any]:
    """Merge results from multiple text chunks"""
    if not results_list:
        return {}
    
    # Initialize with the first chunk's results
    merged = {
        'verb_count': 0,
        'verb_tenses': defaultdict(int),
        'tense_groups': defaultdict(list),
        'passive_count': 0,
        'passive_tenses': defaultdict(int),
        'passive_voice_by_tense_group': defaultdict(list),
        'modal_count': 0,
        'modal_verb_structures': defaultdict(list),
        'all_verbs': [],
        'passive_verbs': [],
        'modal_verbs': []
    }
    
    # Combine all chunks
    for chunk_result in results_list:
        # Simple counters
        merged['verb_count'] += chunk_result['verb_count']
        merged['passive_count'] += chunk_result['passive_count']
        merged['modal_count'] += chunk_result['modal_count']
        
        # Merge dictionaries with integer values
        for tense, count in chunk_result['verb_tenses'].items():
            merged['verb_tenses'][tense] += count
            
        for tense, count in chunk_result['passive_tenses'].items():
            merged['passive_tenses'][tense] += count
        
        # Merge lists in dictionaries
        for tense, verbs in chunk_result['tense_groups'].items():
            merged['tense_groups'][tense].extend(verbs)
            
        for tense, verbs in chunk_result['passive_voice_by_tense_group'].items():
            merged['passive_voice_by_tense_group'][tense].extend(verbs)
            
        for structure, verbs in chunk_result['modal_verb_structures'].items():
            merged['modal_verb_structures'][structure].extend(verbs)
        
        # Merge simple lists
        merged['all_verbs'].extend(chunk_result['all_verbs'])
        merged['passive_verbs'].extend(chunk_result['passive_verbs'])
        merged['modal_verbs'].extend(chunk_result['modal_verbs'])
    
    # Calculate tense group summaries
    tense_group_counts = {tense: len(verbs) for tense, verbs in merged['tense_groups'].items()}
    
    # Calculate modal structure counts
    modal_structure_counts = {
        "modal_base": len(merged['modal_verb_structures'].get("modal_base", [])),
        "modal_passive": len(merged['modal_verb_structures'].get("modal_passive", [])),
        "modal_perfect": len(merged['modal_verb_structures'].get("modal_perfect", []))
    }
    
    # Calculate passive voice by tense group summaries
    passive_tense_group_counts = {tense: len(verbs) for tense, verbs in merged['passive_voice_by_tense_group'].items()}
    
    # Extract just the text from passive_voice_by_tense_group (removing the tense info)
    # This will make it compatible with the template format
    passive_voice_tense_groups = {}
    for tense, verb_tuples in merged['passive_voice_by_tense_group'].items():
        passive_voice_tense_groups[tense] = [verb[0] if isinstance(verb, tuple) else verb for verb in verb_tuples]
    
    # Add calculated counts to the result
    merged['tense_group_counts'] = tense_group_counts
    merged['modal_structure_counts'] = modal_structure_counts
    merged['passive_tense_group_counts'] = passive_tense_group_counts
    merged['passive_voice_by_tense_group_counts'] = passive_tense_group_counts  # Add this line to match the template's expected structure
    merged['passive_voice_tense_groups'] = passive_voice_tense_groups
    
    # Convert defaultdicts to regular dicts for JSON serialization
    for key, value in merged.items():
        if isinstance(value, defaultdict):
            merged[key] = dict(value)
    
    return merged

def analyze_verbs(text: str) -> Dict[str, Any]:
    """Analyze verbs in the given text using parallel processing for better performance"""
    start_time = time.time()
    
    # Chia văn bản thành các đoạn nhỏ để xử lý song song
    chunks = split_text_into_chunks(text)
    logger.info(f"Split text into {len(chunks)} chunks for processing")
    
    # Xử lý song song các đoạn văn bản
    chunk_results = []
    with ProcessPoolExecutor(max_workers=MAX_WORKERS) as executor:
        chunk_results = list(executor.map(process_chunk, chunks))
    
    # Gộp kết quả từ các đoạn
    results = merge_chunk_results(chunk_results)
    
    # Tạo biên bản thông tin thời gian xử lý
    processing_time = time.time() - start_time
    logger.info(f"Verb analysis completed in {processing_time:.2f} seconds")
    results['processing_time'] = f"{processing_time:.2f} seconds"
    
    return results

def determine_tense(token) -> str:
    """Determine the tense of a verb with more specific categorization"""
    tag = token.tag_
    text_lower = token.text.lower()
    sent_text = token.sent.text.lower() if token.sent else ""
    
    # Simple present
    if tag in ["VBP", "VBZ"]:
        # Check for present continuous (am/is/are + verb-ing)
        for i, word in enumerate(token.sent):
            if i > 0 and word.i == token.i and word.tag_ == "VBG" and token.sent[i-1].lemma_ == "be":
                return "present_continuous"
        return "simple_present"
    
    # Simple past
    elif tag == "VBD":
        # Check for past continuous (was/were + verb-ing)
        for i, word in enumerate(token.sent):
            if i > 0 and word.i == token.i and word.tag_ == "VBG" and token.sent[i-1].lemma_ == "be":
                return "past_continuous"
        return "simple_past"
    
    # Base form of verb
    elif tag == "VB":
        # Could be part of future, imperative, or infinitive
        # Check for future tense (will + base form)
        if any(future_aux in sent_text for future_aux in ["will", "shall", "'ll"]):
            return "future_simple"
        # Check for modal + base form
        elif any(token.head.text.lower() == modal for modal in ["can", "could", "may", "might", "must", "should", "would"]):
            return "modal_verb"
        else:
            return "base_form"
    
    # Present participle (used in continuous tenses)
    elif tag == "VBG":
        # Check for present perfect continuous (has/have been + verb-ing)
        if "been" in sent_text and any(aux in sent_text for aux in ["has", "have"]):
            return "present_perfect_continuous"
        # Check for past perfect continuous (had been + verb-ing)
        elif "had been" in sent_text:
            return "past_perfect_continuous"
        # Generic present participle (not caught by specific rules)
        else:
            return "present_participle"
    
    # Past participle (used in perfect tenses and passive voice)
    elif tag == "VBN":
        # Check for present perfect (has/have + past participle)
        if any(aux in sent_text.split() for aux in ["has", "have"]) and all(aux != "been" for aux in sent_text.split()):
            return "present_perfect"
        # Check for past perfect (had + past participle)
        elif "had" in sent_text.split() and "had been" not in sent_text:
            return "past_perfect"
        # Look for passive forms - fix: pass empty dict for aux_verbs parameter
        elif is_passive(token, {}):
            return "passive_voice"
        # Generic past participle
        else:
            return "past_participle"
    
    # Default case for other verb forms
    else:
        return "other"

def is_passive(token, aux_verbs=None) -> bool:
    """Determine if a verb is in passive voice - optimized version"""
    if token.tag_ != "VBN":
        return False
    
    # Nếu aux_verbs đã được tính toán trước, sử dụng nó
    if aux_verbs:
        for ancestor in token.ancestors:
            if ancestor.i in aux_verbs:
                return True
    else:
        # Check if it has an auxiliary verb form of "to be"
        for ancestor in token.ancestors:
            if ancestor.pos_ == "AUX" and ancestor.lemma_ == "be":
                return True
    
    return False

# Thêm background tasks để tránh chặn yêu cầu người dùng
async def analyze_pdf_background(file_path: str, results_callback=None):
    """Analyze PDF in the background and optionally call a callback when done"""
    try:
        # Kiểm tra cache
        cache_path = get_cache_path(file_path)
        if os.path.exists(cache_path):
            logger.info(f"Loading analysis results from cache: {cache_path}")
            results = load_from_cache(cache_path)
            if results_callback:
                results_callback(results)
            return results
            
        # Không có cache, thực hiện phân tích
        text = extract_text_from_pdf(file_path)
        results = analyze_verbs(text)
        
        # Lưu kết quả vào cache
        save_to_cache(cache_path, results)
        
        if results_callback:
            results_callback(results)
        return results
        
    except Exception as e:
        logger.error(f"Error in background PDF analysis: {str(e)}")
        raise

@app.get("/", response_class=HTMLResponse)
async def index(request: Request):
    """Render the home page"""
    return templates.TemplateResponse("index.html", {"request": request})

@app.post("/analyze/", response_class=HTMLResponse)
async def analyze_pdf(request: Request, file: UploadFile = File(...), background_tasks: BackgroundTasks = None):
    """Process uploaded PDF and analyze verbs with optimized performance"""
    logger.info(f"Received file: {file.filename}")
    
    if not file or not file.filename:
        logger.error("No file selected")
        raise HTTPException(status_code=400, detail="No file selected")
    
    if not allowed_file(file.filename):
        logger.error(f"Invalid file type: {file.filename}")
        raise HTTPException(status_code=400, detail="Only PDF files are allowed")
    
    # Create a safe filename 
    safe_filename = os.path.basename(file.filename)
    file_path = os.path.join(UPLOAD_FOLDER, safe_filename)
    
    logger.info(f"Saving file to {file_path}")
    
    # Save the uploaded file
    try:
        # Read file content first
        content = await file.read()
        
        # Then write it to disk
        with open(file_path, "wb") as buffer:
            buffer.write(content)
        
        logger.info(f"File saved successfully: {file_path}")
    except Exception as e:
        logger.error(f"Error saving file: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Error saving file: {str(e)}")
    
    # Extract text first
    try:
        logger.info(f"Starting PDF text extraction")
        extracted_text = extract_text_from_pdf(file_path)
        logger.info(f"Extracted {len(extracted_text)} characters from PDF")
    except Exception as e:
        logger.error(f"Error extracting text from PDF: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Error extracting text from PDF: {str(e)}")
    
    # Kiểm tra nếu kết quả đã được cache
    cache_path = get_cache_path(file_path)
    if os.path.exists(cache_path):
        logger.info(f"Loading analysis results from cache: {cache_path}")
        results = load_from_cache(cache_path)
        logger.info(f"Loaded cached results: Found {results.get('verb_count', 0)} verbs")
        
        # Thêm nội dung văn bản vào results nếu chưa có
        if 'text_content' not in results or not results['text_content']:
            results['text_content'] = extracted_text
            results['extracted_text'] = extracted_text
            # Cập nhật cache với nội dung văn bản
            save_to_cache(cache_path, results)
        
        return templates.TemplateResponse(
            "results.html", 
            {"request": request, "results": results, "filename": safe_filename}
        )
    
    # Extract and analyze text
    try:
        logger.info(f"Starting verb analysis")
        results = analyze_verbs(extracted_text)
        
        # Thêm nội dung văn bản gốc vào kết quả
        results['text_content'] = extracted_text
        results['extracted_text'] = extracted_text
        results['pdf_text'] = extracted_text  # Thêm alias cho template
        
        # Lưu kết quả vào cache để sử dụng sau này
        save_to_cache(cache_path, results)
        
        logger.info(f"Analysis complete: Found {results.get('verb_count', 0)} verbs in {results.get('processing_time', 'N/A')}")
    except Exception as e:
        logger.error(f"Error analyzing text: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Error analyzing text: {str(e)}")
    
    logger.info("Rendering results template")
    return templates.TemplateResponse(
        "results.html", 
        {"request": request, "results": results, "filename": safe_filename}
    )

# API endpoint for JSON response
@app.post("/api/analyze/")
async def analyze_pdf_api(file: UploadFile = File(...), background_tasks: BackgroundTasks = None):
    """Process uploaded PDF and return analysis as JSON with optimized performance"""
    if not file.filename:
        raise HTTPException(status_code=400, detail="No file selected")
    
    if not allowed_file(file.filename):
        raise HTTPException(status_code=400, detail="Only PDF files are allowed")
    
    # Save the uploaded file
    file_path = os.path.join(UPLOAD_FOLDER, file.filename)
    content = await file.read()
    with open(file_path, "wb") as buffer:
        buffer.write(content)
    
    # Extract text first
    extracted_text = extract_text_from_pdf(file_path)
    
    # Kiểm tra cache trước
    cache_path = get_cache_path(file_path)
    if os.path.exists(cache_path):
        logger.info(f"Loading analysis results from cache for API: {cache_path}")
        results = load_from_cache(cache_path)
        
        # Thêm nội dung văn bản nếu chưa có
        if 'text_content' not in results or not results['text_content']:
            results['text_content'] = extracted_text
            results['extracted_text'] = extracted_text
            save_to_cache(cache_path, results)
        
        return results
    
    # Extract and analyze text
    results = analyze_verbs(extracted_text)
    
    # Thêm nội dung văn bản gốc vào kết quả
    results['text_content'] = extracted_text
    results['extracted_text'] = extracted_text
    results['pdf_text'] = extracted_text
    
    # Lưu kết quả vào cache
    save_to_cache(cache_path, results)
    
    return results

# Create export directory if it doesn't exist
if not os.path.exists(EXPORT_FOLDER):
    os.makedirs(EXPORT_FOLDER)

def export_to_pdf(results: Dict[str, Any], filename: str) -> str:
    """Export analysis results to a PDF file using ReportLab"""
    # Create a unique filename for the export
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    output_filename = f"{filename.split('.')[0]}_{timestamp}.pdf"
    output_path = os.path.join(EXPORT_FOLDER, output_filename)
    
    # Create the PDF document
    doc = SimpleDocTemplate(output_path, pagesize=letter)
    styles = getSampleStyleSheet()
    elements = []
    
    # Add title
    title_style = styles["Title"]
    elements.append(Paragraph(f"Verb Analysis Report: {filename}", title_style))
    elements.append(Spacer(1, 12))
    
    # Add timestamp
    date_style = styles["Normal"]
    elements.append(Paragraph(f"Generated on: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}", date_style))
    elements.append(Spacer(1, 12))
    
    # Add summary information
    heading_style = styles["Heading2"]
    normal_style = styles["Normal"]
    
    elements.append(Paragraph("Summary", heading_style))
    elements.append(Paragraph(f"Total Verbs: {results['verb_count']}", normal_style))
    elements.append(Paragraph(f"Passive Voice Verbs: {results['passive_count']}", normal_style))
    elements.append(Paragraph(f"Modal Verbs: {results['modal_count']}", normal_style))
    elements.append(Paragraph(f"Processing Time: {results.get('processing_time', 'N/A')}", normal_style))
    elements.append(Spacer(1, 12))
    
    # Add verb tenses distribution
    elements.append(Paragraph("Verb Tenses Distribution", heading_style))
    tense_data = []
    tense_data.append(["Tense", "Count", "Percentage"])
    
    for tense, count in results['verb_tenses'].items():
        percentage = (count / results['verb_count'] * 100) if results['verb_count'] > 0 else 0
        tense_data.append([tense, str(count), f"{percentage:.1f}%"])
    
    tense_table = Table(tense_data, colWidths=[200, 100, 100])
    tense_table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
        ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
        ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
        ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
        ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
        ('GRID', (0, 0), (-1, -1), 1, colors.black)
    ]))
    elements.append(tense_table)
    elements.append(Spacer(1, 12))
    
    # Add passive voice distribution
    if results['passive_count'] > 0:
        elements.append(Paragraph("Passive Voice Distribution", heading_style))
        passive_data = []
        passive_data.append(["Passive Voice Type", "Count", "Percentage"])
        
        for tense, count in results['passive_tense_group_counts'].items():
            percentage = (count / results['passive_count'] * 100) if results['passive_count'] > 0 else 0
            passive_data.append([tense, str(count), f"{percentage:.1f}%"])
        
        passive_table = Table(passive_data, colWidths=[200, 100, 100])
        passive_table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
            ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
            ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
            ('GRID', (0, 0), (-1, -1), 1, colors.black)
        ]))
        elements.append(passive_table)
        elements.append(Spacer(1, 12))
    
    # Add modal verb structure distribution
    if results['modal_count'] > 0:
        elements.append(Paragraph("Modal Verb Structure Distribution", heading_style))
        modal_data = []
        modal_data.append(["Modal Structure", "Count", "Percentage"])
        
        for structure, count in results['modal_structure_counts'].items():
            percentage = (count / results['modal_count'] * 100) if results['modal_count'] > 0 else 0
            modal_data.append([structure, str(count), f"{percentage:.1f}%"])
        
        modal_table = Table(modal_data, colWidths=[200, 100, 100])
        modal_table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
            ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
            ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
            ('GRID', (0, 0), (-1, -1), 1, colors.black)
        ]))
        elements.append(modal_table)
        elements.append(Spacer(1, 12))
    
    # Build the PDF
    doc.build(elements)
    
    return output_path

def export_to_word(results: Dict[str, Any], filename: str) -> str:
    """Export analysis results to a Word document - fixed data type issues"""
    try:
        from docx import Document
        from docx.shared import Inches, Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from datetime import datetime
        
        logger.info(f"Starting Word export for {filename}")
        logger.info(f"Available keys in results: {list(results.keys())}")
        
        # Create a unique filename for the export
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        output_filename = f"{filename.split('.')[0]}_{timestamp}.docx"
        output_path = os.path.join(EXPORT_FOLDER, output_filename)
        
        # Create a new Document
        document = Document()
        
        # Add title
        title = document.add_heading(f"PDF Verb Analysis Results", level=1)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add filename
        filename_para = document.add_paragraph()
        filename_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        filename_run = filename_para.add_run(f"File: {filename}")
        filename_run.bold = True
        
        # Add timestamp
        timestamp_para = document.add_paragraph()
        timestamp_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        timestamp_run = timestamp_para.add_run(f"Generated on: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        timestamp_run.italic = True
        
        document.add_paragraph()  # Add space
        
        # Add summary section - using safe key access
        document.add_heading("Analysis Summary", level=2)
        
        # Safe key access with defaults and type checking
        verb_count = results.get('verb_count', 0)
        passive_count = results.get('passive_count', 0)
        modal_count = results.get('modal_count', 0)
        processing_time = results.get('processing_time', 'N/A')
        
        # Ensure counts are integers
        try:
            verb_count = int(verb_count) if isinstance(verb_count, (int, float, str)) else 0
            passive_count = int(passive_count) if isinstance(passive_count, (int, float, str)) else 0
            modal_count = int(modal_count) if isinstance(modal_count, (int, float, str)) else 0
        except (ValueError, TypeError):
            verb_count = passive_count = modal_count = 0
        
        summary_points = [
            f"Total Verbs Found: {verb_count}",
            f"Passive Voice Verbs: {passive_count}",
            f"Modal Verbs: {modal_count}",
            f"Processing Time: {processing_time}"
        ]
        
        # Add passive voice percentage if data available
        if verb_count > 0 and passive_count > 0:
            passive_percentage = (passive_count / verb_count * 100)
            summary_points.append(f"Passive Voice Percentage: {passive_percentage:.2f}%")
        
        for point in summary_points:
            para = document.add_paragraph()
            para.style = 'List Bullet'
            para.add_run(point)
        
        document.add_paragraph()  # Add space
        
        # Add verb tense distribution section - if data exists
        tense_counts = results.get('tense_group_counts', {})
        if tense_counts and isinstance(tense_counts, dict):
            document.add_heading("Verb Tense Distribution", level=2)
            
            # Create tense distribution table
            tense_table = document.add_table(rows=1, cols=3)
            tense_table.style = 'Table Grid'
            
            # Add header row
            header_cells = tense_table.rows[0].cells
            header_cells[0].text = "Tense"
            header_cells[1].text = "Count"
            header_cells[2].text = "Percentage"
            
            # Make header bold
            for cell in header_cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True
            
            # Calculate total for percentage - safe calculation
            try:
                # Filter out non-numeric values and ensure we have integers
                valid_counts = []
                for key, value in tense_counts.items():
                    if isinstance(value, (int, float)):
                        valid_counts.append(int(value))
                    elif isinstance(value, str) and value.isdigit():
                        valid_counts.append(int(value))
                    elif isinstance(value, list):
                        valid_counts.append(len(value))  # If it's a list, use its length
                
                total_tense_count = sum(valid_counts) if valid_counts else 0
            except Exception as e:
                logger.warning(f"Error calculating total tense count: {e}")
                total_tense_count = 0
            
            # Add tense data - sorted by count with safe sorting
            try:
                # Create list of (tense, count) pairs with safe count extraction
                tense_pairs = []
                for tense, value in tense_counts.items():
                    if isinstance(value, (int, float)):
                        count = int(value)
                    elif isinstance(value, str) and value.isdigit():
                        count = int(value)
                    elif isinstance(value, list):
                        count = len(value)
                    else:
                        count = 0
                    tense_pairs.append((tense, count))
                
                # Sort by count (descending)
                sorted_tenses = sorted(tense_pairs, key=lambda x: x[1], reverse=True)
                
                for tense, count in sorted_tenses:
                    percentage = (count / total_tense_count * 100) if total_tense_count > 0 else 0
                    row_cells = tense_table.add_row().cells
                    row_cells[0].text = str(tense).replace('_', ' ').title()
                    row_cells[1].text = str(count)
                    row_cells[2].text = f"{percentage:.2f}%"
                    
            except Exception as e:
                logger.error(f"Error processing tense data: {e}")
                # Add a fallback row
                row_cells = tense_table.add_row().cells
                row_cells[0].text = "Error processing data"
                row_cells[1].text = "0"
                row_cells[2].text = "0%"
            
            document.add_paragraph()  # Add space
        
        # Add passive voice section - if data exists
        passive_tenses = results.get('passive_tense_group_counts', {})
        if passive_count > 0 and passive_tenses and isinstance(passive_tenses, dict):
            document.add_heading("Passive Voice Analysis", level=2)
            
            # Add passive voice summary
            passive_percentage = (passive_count / verb_count * 100) if verb_count > 0 else 0
            passive_para = document.add_paragraph()
            passive_para.add_run(f"Total passive voice constructions: {passive_count} ({passive_percentage:.2f}% of all verbs)")
            
            document.add_paragraph()  # Add space
            
            # Create passive voice distribution table
            try:
                passive_table = document.add_table(rows=1, cols=3)
                passive_table.style = 'Table Grid'
                
                # Add header row
                header_cells = passive_table.rows[0].cells
                header_cells[0].text = "Passive Voice Type"
                header_cells[1].text = "Count"
                header_cells[2].text = "Percentage"
                
                # Make header bold
                for cell in header_cells:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.bold = True
                
                # Process passive tenses with safe data handling
                passive_pairs = []
                for tense, value in passive_tenses.items():
                    if isinstance(value, (int, float)):
                        count = int(value)
                    elif isinstance(value, str) and value.isdigit():
                        count = int(value)
                    elif isinstance(value, list):
                        count = len(value)
                    else:
                        count = 0
                    
                    if count > 0:  # Only include tenses with occurrences
                        passive_pairs.append((tense, count))
                
                # Sort by count
                sorted_passive = sorted(passive_pairs, key=lambda x: x[1], reverse=True)
                
                for tense, count in sorted_passive:
                    percentage = (count / passive_count * 100) if passive_count > 0 else 0
                    row_cells = passive_table.add_row().cells
                    row_cells[0].text = str(tense).replace('_', ' ').title()
                    row_cells[1].text = str(count)
                    row_cells[2].text = f"{percentage:.1f}%"
                    
            except Exception as e:
                logger.error(f"Error processing passive voice data: {e}")
            
            document.add_paragraph()  # Add space
        
        # Add modal verb section - if data exists
        modal_structures = results.get('modal_structure_counts', {})
        if modal_count > 0 and modal_structures and isinstance(modal_structures, dict):
            document.add_heading("Modal Verb Analysis", level=2)
            
            # Add modal verb summary
            document.add_paragraph(f"Total modal verbs: {modal_count}")
            
            document.add_paragraph()  # Add space
            
            # Create modal verb structure table
            try:
                modal_table = document.add_table(rows=1, cols=3)
                modal_table.style = 'Table Grid'
                
                # Add header row
                header_cells = modal_table.rows[0].cells
                header_cells[0].text = "Modal Structure"
                header_cells[1].text = "Count"
                header_cells[2].text = "Percentage"
                
                # Make header bold
                for cell in header_cells:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.bold = True
                
                # Process modal structures with safe data handling
                modal_pairs = []
                for structure, value in modal_structures.items():
                    if isinstance(value, (int, float)):
                        count = int(value)
                    elif isinstance(value, str) and value.isdigit():
                        count = int(value)
                    elif isinstance(value, list):
                        count = len(value)
                    else:
                        count = 0
                    
                    if count > 0:  # Only include structures with occurrences
                        modal_pairs.append((structure, count))
                
                # Sort by count
                sorted_modal = sorted(modal_pairs, key=lambda x: x[1], reverse=True)
                
                for structure, count in sorted_modal:
                    percentage = (count / modal_count * 100) if modal_count > 0 else 0
                    row_cells = modal_table.add_row().cells
                    row_cells[0].text = str(structure).replace('_', ' ').title()
                    row_cells[1].text = str(count)
                    row_cells[2].text = f"{percentage:.1f}%"
                    
            except Exception as e:
                logger.error(f"Error processing modal verb data: {e}")
        
        # Add examples section - if data exists
        all_verbs = results.get('all_verbs', [])
        if all_verbs and isinstance(all_verbs, list) and len(all_verbs) > 0:
            document.add_paragraph()  # Add space
            document.add_heading("Sample Verbs", level=2)
            
            try:
                # Limit to first 30 examples for readability
                sample_verbs = all_verbs[:30]
                
                for verb_info in sample_verbs:
                    if isinstance(verb_info, (list, tuple)) and len(verb_info) >= 2:
                        verb, tense = verb_info[0], verb_info[1]
                        document.add_paragraph(f"• {verb} ({str(tense).replace('_', ' ')})", style='ListBullet')
                    elif isinstance(verb_info, str):
                        document.add_paragraph(f"• {verb_info}", style='ListBullet')
                
                if len(all_verbs) > 30:
                    note_para = document.add_paragraph()
                    note_run = note_para.add_run(f"Note: Showing first 30 verbs out of {len(all_verbs)} total verbs found.")
                    note_run.italic = True
                    
            except Exception as e:
                logger.error(f"Error processing verb examples: {e}")
                document.add_paragraph("Error processing verb examples.")
        
        # Add final note
        document.add_paragraph()
        final_note = document.add_paragraph()
        final_note.add_run("This report was generated by PDF Verb Analyzer.").italic = True
        
        # Save the document
        document.save(output_path)
        logger.info(f"Word export successful: {output_path}")
        
        return output_path
        
    except Exception as e:
        logger.error(f"Error exporting to Word: {str(e)}")
        import traceback
        logger.error(f"Full traceback: {traceback.format_exc()}")
        raise Exception(f"Failed to export to Word: {str(e)}")

def export_to_html_pdf(results: Dict[str, Any], filename: str, request: Request) -> str:
    """Export analysis results to PDF using WeasyPrint (HTML to PDF) with working charts"""
    # Create a unique filename for the export
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    output_filename = f"{filename.split('.')[0]}_{timestamp}.pdf"
    output_path = os.path.join(EXPORT_FOLDER, output_filename)
    
    # Ensure text_content is available for highlighting
    if 'text_content' not in results or not results['text_content']:
        # Try other possible keys
        text_content = results.get('extracted_text') or results.get('pdf_text') or ""
        if text_content:
            results['text_content'] = text_content
    
    # Generate HTML content using the template
    html_content = templates.get_template("results.html").render({
        "request": request,
        "results": results,
        "filename": filename,
        "export_view": True,  # Set to True to indicate this is for export
        "include_charts": True  # Flag to explicitly include charts
    })
    
    # Create a temporary HTML file with the rendered content to ensure JS execution
    temp_html_path = os.path.join(tempfile.gettempdir(), f"temp_export_{timestamp}.html")
    
    # Add Chart.js library and necessary scripts to render charts
    chart_script = """
    <script src="https://cdn.jsdelivr.net/npm/chart.js@3.9.1/dist/chart.min.js"></script>
    <script>
    // Function to render charts after page loads
    window.onload = function() {
        // Get tense data from the results
        const tenseLabels = [];
        const tenseData = [];
        
        // Extract data from results
        const tenseGroupCounts = %s;
        
        // Sort data by count (descending)
        const sortedTenses = Object.entries(tenseGroupCounts)
            .sort((a, b) => b[1] - a[1]);
            
        // Get the top 7 tenses for readability
        const topTenses = sortedTenses.slice(0, 7);
        
        // Add an "Other" category for the rest if needed
        if (sortedTenses.length > 7) {
            const otherCount = sortedTenses.slice(7).reduce((sum, item) => sum + item[1], 0);
            topTenses.push(["Other", otherCount]);
        }
        
        // Prepare the data for the chart
        topTenses.forEach(tense => {
            tenseLabels.push(tense[0].replace(/_/g, ' '));
            tenseData.push(tense[1]);
        });
        
        // Create tense distribution chart
        const tenseCtx = document.getElementById('tenseChart').getContext('2d');
        new Chart(tenseCtx, {
            type: 'bar',
            data: {
                labels: tenseLabels,
                datasets: [{
                    label: 'Verb Count',
                    data: tenseData,
                    backgroundColor: [
                        'rgba(75, 192, 192, 0.6)',
                        'rgba(54, 162, 235, 0.6)',
                        'rgba(255, 206, 86, 0.6)',
                        'rgba(255, 99, 132, 0.6)',
                        'rgba(153, 102, 255, 0.6)',
                        'rgba(255, 159, 64, 0.6)',
                        'rgba(199, 199, 199, 0.6)',
                        'rgba(83, 102, 255, 0.6)'
                    ],
                    borderColor: [
                        'rgba(75, 192, 192, 1)',
                        'rgba(54, 162, 235, 1)',
                        'rgba(255, 206, 86, 1)',
                        'rgba(255, 99, 132, 1)',
                        'rgba(153, 102, 255, 1)',
                        'rgba(255, 159, 64, 1)',
                        'rgba(199, 199, 199, 1)',
                        'rgba(83, 102, 255, 1)'
                    ],
                    borderWidth: 1
                }]
            },
            options: {
                scales: {
                    y: {
                        beginAtZero: true
                    }
                },
                plugins: {
                    legend: {
                        display: false
                    }
                },
                animation: false,
                responsive: true,
                maintainAspectRatio: false
            }
        });
        
        // Only create passive voice chart if there are passive verbs
        if (%d > 0) {
            // Create passive voice chart
            const passiveLabels = [];
            const passiveData = [];
            
            // Extract passive voice data
            const passiveCounts = %s;
            
            // Sort and prepare passive data
            const sortedPassive = Object.entries(passiveCounts)
                .filter(item => item[1] > 0)
                .sort((a, b) => b[1] - a[1]);
            
            sortedPassive.forEach(item => {
                passiveLabels.push(item[0].replace(/_/g, ' '));
                passiveData.push(item[1]);
            });
            
            // Only create chart if we have data
            if (passiveLabels.length > 0) {
                const passiveCtx = document.getElementById('passiveChart');
                if (passiveCtx) {
                    new Chart(passiveCtx.getContext('2d'), {
                        type: 'pie',
                        data: {
                            labels: passiveLabels,
                            datasets: [{
                                data: passiveData,
                                backgroundColor: [
                                    'rgba(255, 99, 132, 0.6)',
                                    'rgba(54, 162, 235, 0.6)',
                                    'rgba(255, 206, 86, 0.6)',
                                    'rgba(75, 192, 192, 0.6)',
                                    'rgba(153, 102, 255, 0.6)',
                                    'rgba(255, 159, 64, 0.6)',
                                    'rgba(199, 199, 199, 0.6)'
                                ],
                                borderColor: [
                                    'rgba(255, 99, 132, 1)',
                                    'rgba(54, 162, 235, 1)',
                                    'rgba(255, 206, 86, 1)',
                                    'rgba(75, 192, 192, 1)',
                                    'rgba(153, 102, 255, 1)',
                                    'rgba(255, 159, 64, 1)',
                                    'rgba(199, 199, 199, 1)'
                                ],
                                borderWidth: 1
                            }]
                        },
                        options: {
                            plugins: {
                                legend: {
                                    position: 'right'
                                }
                            },
                            animation: false,
                            responsive: true,
                            maintainAspectRatio: false
                        }
                    });
                }
            }
        }
        
        // Create modal verbs chart if they exist
        if (%d > 0) {
            const modalLabels = [];
            const modalData = [];
            
            // Extract modal verb structure data
            const modalCounts = %s;
            
            // Sort and prepare modal data
            const sortedModal = Object.entries(modalCounts)
                .filter(item => item[1] > 0)
                .sort((a, b) => b[1] - a[1]);
            
            sortedModal.forEach(item => {
                modalLabels.push(item[0].replace(/_/g, ' '));
                modalData.push(item[1]);
            });
            
            // Only create chart if we have data
            if (modalLabels.length > 0) {
                const modalCtx = document.getElementById('modalChart');
                if (modalCtx) {
                    new Chart(modalCtx.getContext('2d'), {
                        type: 'doughnut',
                        data: {
                            labels: modalLabels,
                            datasets: [{
                                data: modalData,
                                backgroundColor: [
                                    'rgba(153, 102, 255, 0.6)',
                                    'rgba(255, 159, 64, 0.6)',
                                    'rgba(75, 192, 192, 0.6)'
                                ],
                                borderColor: [
                                    'rgba(153, 102, 255, 1)',
                                    'rgba(255, 159, 64, 1)',
                                    'rgba(75, 192, 192, 1)'
                                ],
                                borderWidth: 1
                            }]
                        },
                        options: {
                            plugins: {
                                legend: {
                                    position: 'right'
                                }
                            },
                            animation: false,
                            responsive: true,
                            maintainAspectRatio: false
                        }
                    });
                }
            }
        }
        
        // Wait for charts to render, then trigger print
        setTimeout(function() {
            window.print();
        }, 1000);
    };
    </script>
    """
    
    # Insert the chart script with actual data
    tense_group_counts_json = json.dumps(results['tense_group_counts'])
    passive_count = results['passive_count']
    passive_tense_counts_json = json.dumps(results['passive_tense_group_counts'])
    modal_count = results['modal_count']
    modal_structure_counts_json = json.dumps(results.get('modal_structure_counts', {}))
    
    chart_script = chart_script % (tense_group_counts_json, passive_count, passive_tense_counts_json, 
                                  modal_count, modal_structure_counts_json)
    
    # Add chart script and print styles to the HTML content
    html_content = html_content.replace('</head>', f'''
    <style>
        @media print {{
            @page {{
                size: letter portrait;
                margin: 1cm;
            }}
            body {{
                font-size: 10pt;
            }}
            .chart-container {{
                width: 100%;
                height: 300px;
                page-break-inside: avoid;
            }}
            table {{
                page-break-inside: avoid;
            }}
            .section {{
                page-break-inside: avoid;
            }}
            /* Hide UI elements not needed in export */
            .highlight-btn, .btn.small-btn, button, .modal-container, form, nav.navbar {{
                display: none !important;
            }}
            /* Better table styling for print */
            table.data-table {{
                width: 100%;
                border-collapse: collapse;
                margin-bottom: 20px;
            }}
            table.data-table th {{
                background-color: #f2f2f2 !important;
                color: black;
            }}
            table.data-table th, table.data-table td {{
                border: 1px solid #ddd;
                padding: 8px;
            }}
            /* Header and footer */
            .pdf-header {{
                text-align: center;
                margin-bottom: 20px;
            }}
            .pdf-footer {{
                text-align: center;
                font-size: 8pt;
                margin-top: 20px;
                border-top: 1px solid #ddd;
                padding-top: 10px;
            }}
        }}
    </style>
    {chart_script}
    </head>''')
    
    # Add a PDF header with title and date
    pdf_header = f'''
    <div class="pdf-header">
        <h1>Verb Analysis Report: {filename}</h1>
        <p>Generated on: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}</p>
    </div>
    '''
    
    # Add a PDF footer with page number placeholder
    pdf_footer = '''
    <div class="pdf-footer">
        <p>Verb Analysis Report - Page <span class="pageNumber"></span> of <span class="totalPages"></span></p>
    </div>
    '''
    
    # Insert header at the beginning of the body
    html_content = html_content.replace('<body>', f'<body>\n{pdf_header}')
    
    # Insert footer at the end of the body
    html_content = html_content.replace('</body>', f'{pdf_footer}\n</body>')
    
    # Add additional chart canvas for passive voice if needed
    if results['passive_count'] > 0 and '<canvas id="passiveChart"' not in html_content:
        passive_section = '<div class="section">\n            <h2>Passive Voice Distribution</h2>'
        passive_chart_html = passive_section + '''
            <div class="chart-container">
                <canvas id="passiveChart"></canvas>
            </div>'''
        html_content = html_content.replace(passive_section, passive_chart_html)
    
    # Add additional chart canvas for modal verbs if needed
    if results['modal_count'] > 0 and '<canvas id="modalChart"' not in html_content:
        modal_section = '<div class="section">\n            <h2>Modal Verb Distribution</h2>'
        modal_chart_html = modal_section + '''
            <div class="chart-container">
                <canvas id="modalChart"></canvas>
            </div>'''
        
        # Check if the modal section exists, if not add it after the passive section
        if modal_section in html_content:
            html_content = html_content.replace(modal_section, modal_chart_html)
        elif passive_section in html_content:
            # Insert after passive section
            html_content = html_content.replace(passive_section + '''
            <div class="chart-container">
                <canvas id="passiveChart"></canvas>
            </div>''', passive_section + '''
            <div class="chart-container">
                <canvas id="passiveChart"></canvas>
            </div>''' + modal_chart_html)
        else:
            # Insert after tense section
            tense_section = '<div class="section">\n            <h2>Verb Tense Distribution</h2>'
            html_content = html_content.replace(tense_section, tense_section + modal_chart_html)
    
    # Write the complete HTML to a temporary file
    with open(temp_html_path, 'w', encoding='utf-8') as f:
        f.write(html_content)
    
    # Use Puppeteer (via Node.js) to render the page with charts and save as PDF
    # This requires Node.js and Puppeteer to be installed
    try:
        # Try to use Puppeteer for high-quality rendering with charts
        # Fix: Process the file paths before including them in the f-string
        safe_temp_html_path = temp_html_path.replace("\\", "\\\\")
        safe_output_path = output_path.replace("\\", "\\\\")
        
        puppeteer_script = f'''
        const puppeteer = require('puppeteer');
        
        (async () => {{
          const browser = await puppeteer.launch({{
            headless: 'new',
            args: ['--no-sandbox', '--disable-setuid-sandbox']
          }});
          const page = await browser.newPage();
          await page.goto('file://{safe_temp_html_path}', {{waitUntil: 'networkidle0'}});
          
          // Wait for charts to render
          await page.waitForTimeout(2000);
          
          // Add page numbers
          await page.evaluate(() => {{
            const pageNumbers = document.querySelectorAll('.pageNumber');
            pageNumbers.forEach((el, i) => {{
              el.textContent = (i + 1).toString();
            }});
          }});
          
          // Generate PDF
          await page.pdf({{
            path: '{safe_output_path}',
            format: 'Letter',
            printBackground: true,
            margin: {{top: '1cm', right: '1cm', bottom: '1cm', left: '1cm'}},
            displayHeaderFooter: true,
            footerTemplate: '<div style="width: 100%; text-align: center; font-size: 8pt; color: #555;">Page <span class="pageNumber"></span> of <span class="totalPages"></span></div>',
            headerTemplate: '<div></div>'
          }});
          
          await browser.close();
          console.log('PDF generated successfully');
        }})().catch(err => {{
          console.error('Error generating PDF:', err);
          process.exit(1);
        }});
        '''
        
        # Write the Puppeteer script to a temporary file
        script_path = os.path.join(tempfile.gettempdir(), f"puppeteer_script_{timestamp}.js")
        with open(script_path, 'w') as f:
            f.write(puppeteer_script)
        
        # Execute the script with Node.js
        try:
            import subprocess
            result = subprocess.run(['node', script_path], check=True, timeout=30, 
                                  capture_output=True, text=True)
            logger.info(f"Successfully created PDF with charts using Puppeteer: {output_path}")
            logger.debug(f"Puppeteer output: {result.stdout}")
            
            # Check if the PDF was created successfully
            if os.path.exists(output_path) and os.path.getsize(output_path) > 0:
                # Clean up temporary files
                os.remove(temp_html_path)
                os.remove(script_path)
                return output_path
            else:
                logger.warning("Puppeteer did not generate a valid PDF. Falling back to WeasyPrint")
        except (subprocess.SubprocessError, FileNotFoundError) as e:
            logger.warning(f"Puppeteer PDF generation failed: {e}. Falling back to WeasyPrint")
            # Fall back to WeasyPrint if Puppeteer fails
    except Exception as e:
        logger.warning(f"Error setting up Puppeteer: {e}. Falling back to WeasyPrint")
    
    # Fallback: Use WeasyPrint (which won't render the charts properly, but will provide the tables)
    try:
        # Add a message to indicate charts are not available in this version
        html_content = html_content.replace('<div class="chart-container">', 
                                          '<div class="chart-container"><p style="text-align:center; color:#777;">Charts not available in this export method</p>')
        
        # WeasyPrint fallback (doesn't render JavaScript)
        HTML(string=html_content).write_pdf(
            output_path,
            stylesheets=[CSS(string='''
                @page {
                    size: letter portrait;
                    margin: 1cm;
                }
                body {
                    font-size: 10pt;
                }
                .chart-container {
                    border: 1px dashed #ccc;
                    padding: 10px;
                    margin: 10px 0;
                    text-align: center;
                }
                table {
                    page-break-inside: avoid;
                    width: 100%;
                    border-collapse: collapse;
                    margin-bottom: 20px;
                }
                table th {
                    background-color: #f2f2f2;
                }
                table th, table td {
                    border: 1px solid #ddd;
                    padding: 8px;
                }
                .section {
                    page-break-inside: avoid;
                    margin-bottom: 20px;
                }
                /* Hide UI elements not needed in export */
                .highlight-btn, .btn.small-btn, button, .modal-container, form, nav.navbar {
                    display: none !important;
                }
                /* Header and footer */
                .pdf-header {
                    text-align: center;
                    margin-bottom: 20px;
                }
                .pdf-footer {
                    text-align: center;
                    font-size: 8pt;
                    margin-top: 20px;
                    border-top: 1px solid #ddd;
                    padding-top: 10px;
                }
            ''')]
        )
        
        # Clean up temporary file
        os.remove(temp_html_path)
        
        logger.info(f"Created PDF with WeasyPrint (fallback method): {output_path}")
        return output_path
        
    except Exception as e:
        logger.error(f"Error creating PDF with WeasyPrint: {e}")
        raise

@app.post("/export/")
async def export_results(request: Request, filename: str = Form(...), format: str = Form(...)):
    """Export analysis results to the selected format (PDF or Word)"""
    file_path = os.path.join(UPLOAD_FOLDER, filename)
    
    if not os.path.exists(file_path):
        raise HTTPException(status_code=404, detail="File not found")
    
    # Get cached analysis results
    cache_path = get_cache_path(file_path)
    if not os.path.exists(cache_path):
        raise HTTPException(status_code=404, detail="Analysis results not found")
    
    results = load_from_cache(cache_path)
    
    try:
        if format == 'pdf':
            output_path = export_to_pdf(results, filename)
            return FileResponse(
                path=output_path,
                filename=os.path.basename(output_path),
                media_type="application/pdf"
            )
        elif format == 'docx':
            output_path = export_to_word(results, filename)
            return FileResponse(
                path=output_path,
                filename=os.path.basename(output_path),
                media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
        elif format == 'html_pdf':
            output_path = export_to_html_pdf(results, filename, request)
            return FileResponse(
                path=output_path,
                filename=os.path.basename(output_path),
                media_type="application/pdf"
            )
        else:
            raise HTTPException(status_code=400, detail="Invalid export format")
    except Exception as e:
        logger.error(f"Error exporting results: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Error exporting results: {str(e)}")

@app.get("/export/pdf/{filename}")
async def export_pdf(request: Request, filename: str):
    """Export analysis results to PDF format"""
    file_path = os.path.join(UPLOAD_FOLDER, filename)
    
    if not os.path.exists(file_path):
        raise HTTPException(status_code=404, detail="File not found")
    
    # Get cached analysis results
    cache_path = get_cache_path(file_path)
    if not os.path.exists(cache_path):
        raise HTTPException(status_code=404, detail="Analysis results not found")
    
    results = load_from_cache(cache_path)
    
    try:
        output_path = export_to_pdf(results, filename)
        return FileResponse(
            path=output_path,
            filename=os.path.basename(output_path),
            media_type="application/pdf"
        )
    except Exception as e:
        logger.error(f"Error exporting to PDF: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Error exporting to PDF: {str(e)}")

@app.get("/export/docx/{filename}")
async def export_docx(request: Request, filename: str):
    """Export analysis results to Word document format"""
    file_path = os.path.join(UPLOAD_FOLDER, filename)
    
    if not os.path.exists(file_path):
        raise HTTPException(status_code=404, detail="File not found")

    # Get cached analysis results
    cache_path = get_cache_path(file_path)
    if not os.path.exists(cache_path):
        raise HTTPException(status_code=404, detail="Analysis results not found")

    results = load_from_cache(cache_path)

    try:
        output_path = export_to_word(results, filename)
        return FileResponse(
            path=output_path,
            filename=os.path.basename(output_path),
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
    except Exception as e:
        logger.error(f"Error exporting to Word: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Error exporting to Word: {str(e)}")
    
if __name__ == "__main__":
    import uvicorn
    uvicorn.run(
        "app:app", 
        host="0.0.0.0", 
        port=PORT,
        reload=False  # Disable reload in production
    )